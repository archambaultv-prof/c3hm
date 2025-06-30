from importlib import resources
from pathlib import Path

import docx
import docx.enum.text
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.table import Table

from c3hm.data.evaluation.criterion import Criterion
from c3hm.data.gradesheet import GradeSheet
from c3hm.data.rubric.rubric import Rubric
from c3hm.data.student.student import Student
from c3hm.utils.word import (
    add_word_table,
    set_as_table_header,
    set_cell_background,
    set_orientation,
    set_row_borders,
)

PERFECT_GREEN    = "C8FFC8"  # RGB(200, 255, 200)
VERY_GOOD_GREEN  = "F0FFB0"  # RGB(240, 255, 176)
HALF_WAY_YELLOW  = "FFF8C2"  # RGB(255, 248, 194)
MINIMAL_RED      = "FFE4C8"  # RGB(255, 228, 200)
BAD_RED          = "FFC8C8"  # RGB(255, 200, 200)

def scale_color_schemes(size: int) -> list[str] | None:
    """
    Retourne une liste de couleurs format Word pour les cellules liées au barème.

    La première couleur est le vert parfait, la dernière est le rouge insuffisant.
    Prend en charge 2 à 5 niveaux de barème. Autrement, renvoie None.
    """
    color_schemes = {
        2: [PERFECT_GREEN, BAD_RED],
        3: [PERFECT_GREEN, HALF_WAY_YELLOW, BAD_RED],
        4: [PERFECT_GREEN, VERY_GOOD_GREEN, HALF_WAY_YELLOW, BAD_RED],
        5: [PERFECT_GREEN, VERY_GOOD_GREEN, HALF_WAY_YELLOW, MINIMAL_RED, BAD_RED],
    }
    return color_schemes.get(size)


def generate_rubric(
        rubric: Rubric,
        output_path: Path | str,
        grade_sheet: GradeSheet | None = None,
        student: Student | None = None,
        ) -> None:
    """
    Génère un document Word pour les étudiants à partir d’un grille d’évaluation (Rubric).
    """
    if not grade_sheet and student:
        raise ValueError("Une feuille de notes doit être fournie si un étudiant est défini.")
    if grade_sheet and not student:
        raise ValueError("Un étudiant doit être fourni si une feuille de notes est définie.")

    output_path = Path(output_path)
    eval = rubric.evaluation

    # forcer l'extension .docx
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    # démarrer Word avec le modèle par défaut
    # Ce dernier contient déjà le formatage et les styles nécessaires
    template = resources.files("c3hm.assets.templates.word").joinpath("rubric-default.docx")
    with resources.as_file(template) as path:
        doc = docx.Document(str(path))

    # Modifier l'orientation de la page au besoin
    if rubric.format.orientation:
        set_orientation(doc, rubric.format.orientation)
    elif len(rubric.grade_levels.levels) >= 4:
        set_orientation(doc, "paysage")
    else:
        set_orientation(doc, "portrait")

    # insérer le titre
    p = doc.paragraphs[0]
    p.text = eval.title()
    p.style = "Heading 1"
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    # insérer le nom de l'étudiant
    if grade_sheet and student:
        p = doc.add_paragraph()
        p.text = f"{student.first_name} {student.last_name}"

    # Insérer le commentaire général
    if grade_sheet and grade_sheet.has_comment(eval):
        p = doc.add_paragraph()
        p.text = f"Commentaire : {grade_sheet.get_comment(eval)}"
        p.style = "Normal"
        p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

    # insérer la table pour la grille
    if grade_sheet and grade_sheet.has_criteria_comments(eval):
        col_width = rubric.format.columns_width_comments
        nb_columns = len(rubric.grade_levels.levels) + 2
    else:
        col_width = rubric.format.columns_width
        nb_columns = len(rubric.grade_levels.levels) + 1
    table = add_word_table(doc, nb_columns, col_width)

    # Remplir la première ligne avec le barème
    set_first_row(rubric, table, grade_sheet)

    # Remplir le reste avec critères et indicateurs dans l'ordre
    for criterion in eval.criteria:
        add_criterion(table, criterion, rubric, grade_sheet)

    # Ajouter une bordure en bas de la dernière ligne
    last_row = table.rows[-1]
    set_row_borders(last_row, top=None, bottom=1.0)

    # enregistrer le fichier
    doc.save(str(output_path))


def add_criterion(table: Table,
                  criterion: Criterion,
                  rubric: Rubric,
                  grade_sheet: GradeSheet | None = None) -> None:
    """
    Ajoute un critère et ses indicateurs à la table.
    """
    row = table.add_row()
    narrow_nbsp = "\u202F"

    # Critère
    p = row.cells[0].paragraphs[0]
    txt = criterion.name
    pts = "pt" if criterion.points == 1 else "pts"
    txt += f" ({criterion.points}{narrow_nbsp}{pts})"
    p.text = txt
    p.style = "Heading 3"

    if grade_sheet:
        c_grade = grade_sheet.get_grade(criterion, rubric.evaluation.precision + 1)
        c_percentage = grade_sheet.get_percentage(criterion)
        c_grade_pos = rubric.grade_levels.get_index_by_percentage(c_percentage)
        p = row.cells[c_grade_pos + 1].paragraphs[0]
        pts = "pt" if c_grade == 1 else "pts"
        p.text = f"{c_grade}{narrow_nbsp}{pts}"
        p.style = "Heading 3"

        # Ajoute le commentaire du critère si disponible
        if grade_sheet.has_comment(criterion):
            comment_cell = row.cells[-1]
            p = comment_cell.paragraphs[0]
            p.text = grade_sheet.get_comment(criterion)

    # Indicateurs
    for i, indicator in enumerate(criterion.indicators):
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        run = p.add_run()
        txt = indicator.name
        if rubric.format.show_indicators_points:
            # Afficher le pourcentage de l'indicateur
            pts = "pt" if indicator.points == 1 else "pts"
            txt += f" ({indicator.points}{narrow_nbsp}{pts})"
        run.text = txt
        run.style = "Emphasis"

        if grade_sheet and grade_sheet.has_comment(indicator):
            comment_cell = row.cells[-1]
            p = comment_cell.paragraphs[0]
            p.text = grade_sheet.get_comment(indicator)

        # Si l'indicateur est noté, on trouve le niveau de note
        if grade_sheet:
            i_percentage = grade_sheet.get_percentage(indicator)
            i_grade_pos = rubric.grade_levels.get_index_by_percentage(i_percentage)
        else:
            i_grade_pos = None

        for i, level in enumerate(rubric.grade_levels.levels):
            # Descripteur
            cell = row.cells[i + 1]
            cell.text = rubric.descriptors.get_descriptor(indicator, level)

            # Ajoute la note de l'indicateur si disponible
            if grade_sheet and i == i_grade_pos:
                if rubric.format.show_indicators_points:
                    g = grade_sheet.get_grade(indicator, rubric.evaluation.precision + 1)
                    run = cell.paragraphs[0].add_run(f" ({g})")
                    run.style = "Emphasis"
                # On met en surbrillance le descripteur
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_first_row(rubric: Rubric, table: Table, grade_sheet: GradeSheet | None = None):
    """
    Remplit la première ligne du tableau avec le barème et les seuils.
    """
    eval = rubric.evaluation

    # En-tête de la table
    set_as_table_header(table.rows[0])
    set_row_borders(table.rows[0])

    # Première cellule : Total sur X pts
    hdr_cells = table.rows[0].cells
    total_cell = hdr_cells[0]
    p = total_cell.paragraphs[0]
    if grade_sheet:
        total = grade_sheet.get_grade(eval, eval.precision)
        points = eval.points
        t = f"Note : {total}"
        t += f" / {points}"
        p.text = t
        p.style = "Heading 3"

    # Barème et seuils
    set_grade_levels(rubric, hdr_cells)

    # Commentaires
    if grade_sheet and grade_sheet.has_criteria_comments(eval):
        comment_cell = hdr_cells[-1]
        p = comment_cell.paragraphs[0]
        p.text = "Commentaire"
        p.style = "Heading 3"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_grade_levels(rubric: Rubric, hdr_cells):
    color_schemes = scale_color_schemes(len(rubric.grade_levels.levels))
    for i, label in enumerate(rubric.grade_levels.levels):
        cell = hdr_cells[i + 1]
        if color_schemes:
            set_cell_background(cell, color_schemes[i])
        if rubric.format.show_level_descriptions:
            t_str = label.level_description()
            txt = f"{label.name}\n({t_str})"
        else:
            txt = label.name
        p = cell.paragraphs[0]
        p.text = txt
        p.style = "Heading 3"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
