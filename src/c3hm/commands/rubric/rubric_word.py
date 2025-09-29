from decimal import Decimal
from importlib import resources
from pathlib import Path

import docx
import docx.enum.text
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph

from c3hm.data.config import Config
from c3hm.data.criterion import Criterion
from c3hm.utils.word import (
    add_word_table,
    set_as_table_header,
    set_cell_background,
    set_orientation,
    set_row_borders,
)

PERFECT_GREEN    = "C8FFC8"  # RGB(200, 255, 200)
VERY_GOOD_GREEN  = "F0FFB0"  # RGB(240, 255, 176)
HALF_WAY_YELLOW  = "FFF8C2"  # RGB(255, 248, 194)
MINIMAL_RED      = "FFE4C8"  # RGB(255, 228, 200)
BAD_RED          = "FFC8C8"  # RGB(255, 200, 200)

def scale_color_schemes(size: int) -> list[str] | None:
    """
    Retourne une liste de couleurs format Word pour les cellules liées au barème.

    La première couleur est le vert parfait, la dernière est le rouge insuffisant.
    Prend en charge 2 à 5 niveaux de barème. Autrement, renvoie None.
    """
    color_schemes = {
        2: [PERFECT_GREEN, BAD_RED],
        3: [PERFECT_GREEN, HALF_WAY_YELLOW, BAD_RED],
        4: [PERFECT_GREEN, VERY_GOOD_GREEN, HALF_WAY_YELLOW, BAD_RED],
        5: [PERFECT_GREEN, VERY_GOOD_GREEN, HALF_WAY_YELLOW, MINIMAL_RED, BAD_RED],
    }
    return color_schemes.get(size)


def generate_rubric_word(
        config: Config,
        output_path: Path | str,
        ) -> None:
    """
    Génère un document Word pour les étudiants à partir d’un grille d’évaluation (Rubric).
    """
    output_path = Path(output_path)

    # forcer l'extension .docx
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    # démarrer Word avec le modèle par défaut
    # Ce dernier contient déjà le formatage et les styles nécessaires
    template = resources.files("c3hm.assets.templates.word").joinpath("rubric-default.docx")
    with resources.as_file(template) as path:
        doc = docx.Document(str(path))

    # Modifier l'orientation de la page au besoin
    if config.format.orientation:
        set_orientation(doc, config.format.orientation)
    elif len(config.grade_levels) >= 4:
        set_orientation(doc, "paysage")
    else:
        set_orientation(doc, "portrait")

    # insérer le titre
    p = doc.paragraphs[0]
    p.text = config.evaluation_title()
    p.style = "Heading 1"
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    # insérer le nom de l'étudiant
    if config.is_graded:
        p = doc.add_paragraph()
        add_student_name(p, config.get_student_full_name())

        # Insérer le commentaire général
        if config.evaluation_comment:
            p = doc.add_paragraph()
            add_comment_runs(p, config.evaluation_comment)

    # insérer la table pour la grille
    col_width = config.format.columns_width
    nb_columns = len(config.grade_levels) + 1
    table = add_word_table(doc, nb_columns, col_width)

    # Remplir la première ligne avec le barème
    set_first_row(config, table)

    # Remplir le reste avec critères et indicateurs dans l'ordre
    for criterion in config.criteria:
        try:
            add_criterion(table, criterion, config)
        except Exception as e:
            raise ValueError(f"Erreur lors de l'ajout du critère '{criterion.name}' "
                             f"pour l'étudiant '{config.get_student_full_name()}'.") from e

    # Ajouter une bordure en bas de la dernière ligne
    last_row = table.rows[-1]
    set_row_borders(last_row, top=None, bottom=1.0)

    # enregistrer le fichier
    doc.save(str(output_path))


def add_criterion(table: Table,
                  criterion: Criterion,
                  config: Config) -> None:
    """
    Ajoute un critère et ses indicateurs à la table.
    """
    row = table.add_row()
    narrow_nbsp = "\u202F"

    # Critère
    p = row.cells[0].paragraphs[0]
    txt = criterion.name
    pts = "pt" if criterion.total == 1 else "pts"
    txt += f" ({criterion.total}{narrow_nbsp}{pts})"
    p.text = txt
    p.style = "Heading 3"

    if config.is_graded:
        c_grade = criterion.get_grade()
        c_percentage = c_grade / criterion.get_total()
        c_grade_pos = config.get_level_index_by_percentage(c_percentage)
        p = row.cells[c_grade_pos + 1].paragraphs[0]
        pts = "pt" if c_grade == 1 else "pts"
        c_percent = f"({c_percentage * 100:.0f}%)"
        p.text = f"{c_grade}{narrow_nbsp}{pts} {c_percent}"
        p.style = "Heading 3"

        # Ajoute le commentaire du critère si disponible
        if criterion.comment:
            add_comment_row(table, criterion.comment)

    # Indicateurs
    for i, indicator in enumerate(criterion.indicators):
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        run = p.add_run()
        txt = indicator.name
        if config.format.show_indicators_points:
            # Afficher le pourcentage de l'indicateur
            pts = "pt" if indicator.points == 1 else "pts"
            txt += f" ({indicator.points}{narrow_nbsp}{pts})"
        run.text = txt
        run.style = "Emphasis"

        # Si l'indicateur est noté, on trouve le niveau de note
        if indicator.grade:
            i_percentage = indicator.get_grade() / indicator.get_total()
            i_grade_pos = config.get_level_index_by_percentage(i_percentage)
            if indicator.comment:
                add_comment_row(table, indicator.comment)
        else:
            i_grade_pos = None

        for i, descriptor in enumerate(indicator.descriptors):
            # Descripteur
            cell = row.cells[i + 1]
            if descriptor is not None:
                cell.text = descriptor
            else:
                cell.text = "N/A"
                if indicator.grade and i == i_grade_pos:
                    raise ValueError(f"Note attribuée à l'indicateur {indicator.name} "
                                     f"pour l'eleve {config.get_student_full_name()}, "
                                     f" mais il n'y a pas de descripteur défini pour ce niveau.")

            if indicator.grade and i == i_grade_pos:
                # Ajoute la note de l'indicateur si disponible
                # et si demandé dans la config
                if config.format.show_indicators_points:
                    g = indicator.get_grade()
                    run = cell.paragraphs[0].add_run(f" ({g})")
                    run.style = "Emphasis"
                # On met en surbrillance le descripteur
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def add_comment_row(table: Table, comment: str) -> None:
    comment_row = table.add_row()
    # Merge from second cell to the end on the comment row
    merged_cell = comment_row.cells[1].merge(comment_row.cells[-1])
    p = merged_cell.paragraphs[0]
    # Make the label bold, followed by the comment
    add_comment_runs(p, comment)

def add_student_name(p: Paragraph, student: str) -> None:
    label_run = p.add_run("Étudiant : ")
    label_run.bold = True
    label_run.style = "Emphasis"
    name_run = p.add_run(student)
    name_run.style = "Emphasis"

def add_comment_runs(p: Paragraph, comment: str) -> None:
    label_run = p.add_run("Commentaire : ")
    label_run.bold = True
    label_run.style = "Emphasis"
    # Apply the character style to the comment run (not the paragraph)
    comment_run = p.add_run(comment)
    comment_run.style = "Emphasis"

def set_first_row(config: Config, table: Table):
    """
    Remplit la première ligne du tableau avec le barème et les seuils.
    """

    # En-tête de la table
    set_as_table_header(table.rows[0])
    set_row_borders(table.rows[0], top=1.0, bottom=0.5)

    # Première cellule : Total sur X pts
    hdr_cells = table.rows[0].cells
    total_cell = hdr_cells[0]
    p = total_cell.paragraphs[0]
    points = config.evaluation_total
    ndigits = config.evaluation_total_nb_decimals
    if config.is_graded:
        total = config.evaluation_grade
        t = f"Note : {total:.{ndigits}f} / {points:.{ndigits}f}"
    else:
        pts = "pt" if points == Decimal("1") else "pts"
        t = f"Total sur {points:.{ndigits}f} {pts}"
    p.text = t
    p.style = "Heading 3"

    # Barème et seuils
    set_grade_levels(config, hdr_cells)


def set_grade_levels(config: Config, hdr_cells):
    color_schemes = scale_color_schemes(len(config.grade_levels))
    for i, label in enumerate(config.grade_levels):
        cell = hdr_cells[i + 1]
        if color_schemes:
            set_cell_background(cell, color_schemes[i])
        if config.format.show_grade_level_descriptions:
            t_str = label.level_description()
            txt = f"{label.name}\n({t_str})"
        else:
            txt = label.name
        p = cell.paragraphs[0]
        p.text = txt
        p.style = "Heading 3"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
