import sys
from importlib import resources
from pathlib import Path
from typing import Any

import docx
import docx.enum.text
from docx.document import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from docx.table import Table

from c3hm.data.rubric import (
    CTHM_GLOBAL_COMMENT,
    Criterion,
    Rubric,
)
from c3hm.data.student import Student
from c3hm.utils.decimal import decimal_to_number, round_to_nearest_quantum

PERFECT_GREEN    = "C8FFC8"  # RGB(200, 255, 200)
VERY_GOOD_GREEN  = "F0FFB0"  # RGB(240, 255, 176)
HALF_WAY_YELLOW  = "FFF8C2"  # RGB(255, 248, 194)
MINIMAL_RED      = "FFE4C8"  # RGB(255, 228, 200)
BAD_RED          = "FFC8C8"  # RGB(255, 200, 200)

def scale_color_schemes(size: int) -> list[str] | None:
    """
    Retourne une liste de couleurs format Word pour les cellules liées au barème.

    La première couleur est le vert parfait, la dernière est le rouge insuffisant.
    Prend en charge 2 à 5 niveaux de barème. Autrement, renvoie None.
    """
    color_schemes = {
        2: [PERFECT_GREEN, BAD_RED],
        3: [PERFECT_GREEN, HALF_WAY_YELLOW, BAD_RED],
        4: [PERFECT_GREEN, VERY_GOOD_GREEN, HALF_WAY_YELLOW, BAD_RED],
        5: [PERFECT_GREEN, VERY_GOOD_GREEN, HALF_WAY_YELLOW, MINIMAL_RED, BAD_RED],
    }
    return color_schemes.get(size)

def set_as_table_header(row):
    """
    Définit la ligne d'un tableau Word comme en-tête répétée sur chaque page.
    """
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()

    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def set_cell_background(cell, color_hex: str):
    """
    Définit la couleur de fond d'une cellule dans un tableau Word.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')         # shading style
    shd.set(qn('w:color'), 'auto')        # text color (auto = black/white depending on bg)
    shd.set(qn('w:fill'), color_hex)      # background color

    # remove old <w:shd> if any
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    tc_pr.append(shd)

def set_row_borders(row, top: float | None =0.5, bottom: float | None =0.5):
    """
    Définit les bordures hautes et basses d'une ligne dans un tableau Word.
    """
    def set_border(tc_borders, side, width_pt):
        border = tc_borders.find(qn(f'w:{side}'))
        if border is None:
            border = OxmlElement(f'w:{side}')
            tc_borders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(int(width_pt * 8)))  # Word uses 1/8 pt units
        border.set(qn('w:space'), "0")

    for cell in row.cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        if top:
            set_border(tc_borders, 'top', top)
        if bottom:
            set_border(tc_borders, 'bottom', bottom)

def set_orientation(doc: Document, orientation: str):
    """
    Définit l'orientation de la page dans un document Word.
    """
    sections = doc.sections
    for section in sections:
        if orientation == "paysage":
            section.orientation = WD_ORIENT.LANDSCAPE
            # Échanger la largeur et la hauteur de la page
            if section.page_width < section.page_height: # type: ignore
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )
        elif orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
            # Échanger la largeur et la hauteur de la page
            if section.page_width > section.page_height: # type: ignore
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )
        else:
            raise ValueError(f"Orientation '{orientation}' non valide. "
                             "Utilisez 'paysage' ou 'portrait'.")


def generate_rubric(
        rubric: Rubric,
        output_path: Path | str,
        *,
        title: str = "Grille d'évaluation",
        student: Student | None = None,
        grades: dict[str, Any] | None = None,
        ) -> None:
    """
    Génère un document Word pour les étudiants à partir d’un grille d’évaluation (Rubric).
    """
    output_path = Path(output_path)

    # forcer l'extension .docx
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    # démarrer Word avec le modèle par défaut
    # Ce dernier contient déjà le formatage et les styles nécessaires
    template = resources.files("c3hm.assets.templates.word").joinpath("rubric-default.docx")
    with resources.as_file(template) as path:
        doc = docx.Document(str(path))

    # Modifier l'orientation de la page au besoin
    if rubric.format.orientation:
        set_orientation(doc, rubric.format.orientation)
    elif rubric.nb_criteria() >= 4:
        set_orientation(doc, "landscape")
    else:
        set_orientation(doc, "portrait")

    # insérer le titre
    p = doc.paragraphs[0]
    p.text = title
    p.style = "Heading 1"
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    # insérer le nom de l'étudiant
    if grades and student:
        p = doc.add_paragraph()
        p.text = f"{student.first_name} {student.last_name}"

    # Insérer le commentaire général
    if grades and CTHM_GLOBAL_COMMENT in grades:
        comment = str(grades[CTHM_GLOBAL_COMMENT]).strip()
        if comment:
            p = doc.add_paragraph()
            p.text = f"Commentaire : {comment}"
            p.style = "Normal"
            p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

    # insérer la table pour la grille
    has_criteria_comments = any(
        criterion_has_comments(criterion, grades)
        for criterion in rubric.criteria
    )
    if grades and has_criteria_comments:
        col_width = rubric.format.columns_width_comments
        nb_columns = rubric.nb_criteria() + 2
    else:
        col_width = rubric.format.columns_width
        nb_columns = rubric.nb_criteria() + 1
    table = add_word_table(doc, nb_columns, col_width)

    # Remplir la première ligne avec le barème
    set_first_row(rubric, table, grades)

    # Remplir le reste avec critères et indicateurs dans l'ordre
    for criterion in rubric.criteria:
        add_criterion(table, criterion, rubric, grades)

    # Ajouter une bordure en bas de la dernière ligne
    last_row = table.rows[-1]
    set_row_borders(last_row, top=None, bottom=1.0)

    # enregistrer le fichier
    doc.save(str(output_path))

def add_word_table(doc: Document, n_cols: int, column_widths_cm: list[float|None]):
    """
    Insère un tableau à n_cols colonnes, en utilisant largeurs_cm pour définir
    la largeur de chaque colonne en cm. Les valeurs None partagent
    équitablement l’espace restant.
    """
    # 1) Création du tableau
    table = doc.add_table(rows=1, cols=n_cols, style="Normal Table")
    if not column_widths_cm:
        return table

    table.autofit = False  # désactive l’ajustement automatique de Word

    # 2) Calcul de la largeur disponible en EMU
    sect = doc.sections[0]
    avail_width_emu = (
        sect.page_width
        - sect.left_margin # type: ignore
        - sect.right_margin
    )

    # 3) Conversion des spécifications cm en EMU
    fixed_emu = [Cm(w) for w in column_widths_cm if w is not None]
    total_fixed = sum(fixed_emu)

    # Décompte du nombre de colonnes à largeur automatique
    nb_auto = sum(1 for w in column_widths_cm if w is None)
    emu_auto = (
        (avail_width_emu - total_fixed) // nb_auto
        if nb_auto else
        0
    )

    # 4) Assignation de la largeur à chaque colonne
    for idx, col in enumerate(table.columns):
        w = column_widths_cm[idx]
        col.width = Cm(w) if w is not None else emu_auto # type: ignore

    # 5) Assignation de la largeur à chaque cellule
    # (yep ... c'est comme ça dans Word)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            w = column_widths_cm[idx]
            cell.width = Cm(w) if w is not None else emu_auto # type: ignore

    return table


def add_criterion(table: Table,
                  criterion: Criterion,
                  rubric: Rubric,
                  grades: dict[str, Any] | None = None):
    """
    Ajoute un critère et ses indicateurs à la table.
    """
    row = table.add_row()
    percent_gray = RGBColor(0x80, 0x80, 0x80)
    narrow_nbsp = "\u202F"

    # Critère
    p = row.cells[0].paragraphs[0]
    p.text = criterion.name.strip()
    p.style = "Heading 3"
    run = p.add_run(f" ({decimal_to_number(criterion.percentage)}{narrow_nbsp}%)") # type: ignore
    run.style = "Heading 3"
    run.font.color.rgb = percent_gray

    if grades:
        if grades[criterion.xl_grade_overwrite_cell_id()] is not None:
            # Si la note est écrasée manuellement, on l'utilise
            c_grade = grades[criterion.xl_grade_overwrite_cell_id()]
        else:
            # Sinon, on calcule la note du critère
            c_grade = sum(
                grades[ind.xl_grade_cell_id()] * ind.percentage
                for ind in criterion.indicators
            ) / criterion.percentage # type: ignore
            c_grade = round_to_nearest_quantum(c_grade, rubric.precision)
        # Find the position according to the grades_thresholds
        grade_pos = len(rubric.grade_thresholds)  # Default to last position
        for i, (_, min_grade, _) in enumerate(rubric.grade_thresholds):
            if c_grade >= min_grade:
                grade_pos = i + 1
                break
        p = row.cells[grade_pos].paragraphs[0]
        p.text = f"{decimal_to_number(c_grade)}"
        p.style = "Heading 3"

    # Indicateurs
    for indicator in criterion.indicators:
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        run = p.add_run(indicator.name)
        run.style = "Emphasis"
        if rubric.format.show_indicators_percent:
            # Afficher le pourcentage de l'indicateur
            run = p.add_run(f" ({decimal_to_number(indicator.percentage)}{narrow_nbsp}%)") # type: ignore
            run.style = "Emphasis"
            run.font.color.rgb = percent_gray

        # Descripteurs alignés avec les niveaux de barème
        if grades:
            i_grade = grades[indicator.xl_grade_cell_id()]
            grade_pos = len(rubric.grade_thresholds)  # Default to last position
            for i, (_, min_grade, _) in enumerate(rubric.grade_thresholds):
                if i_grade >= min_grade:
                    grade_pos = i + 1
                    break
        else:
            grade_pos = None

        color_schemes = scale_color_schemes(len(rubric.grade_levels))
        for i, descriptor in enumerate(indicator.descriptors):
            # Descripteur
            cell = row.cells[i + 1]
            cell.text = descriptor

            # Ajoute la note de l'indicateur si disponible
            if grades and i+1 == grade_pos:
                i_grade = grades[indicator.xl_grade_cell_id()]
                run = cell.paragraphs[0].add_run(f" ({decimal_to_number(i_grade)})")
                run.style = "Emphasis"
                run.font.color.rgb = percent_gray
                if color_schemes:
                    set_cell_background(cell, color_schemes[i])
                else:
                    # On met en surbrillance le descripteur
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def set_first_row(rubric: Rubric, table: Table, grades: dict[str, Any] | None):
    """
    Remplit la première ligne du tableau avec le barème et les seuils.
    """

    # En-tête de la table
    set_as_table_header(table.rows[0])
    set_row_borders(table.rows[0])

    # Première cellule : Total sur X pts
    hdr_cells = table.rows[0].cells
    total_cell = hdr_cells[0]
    p = total_cell.paragraphs[0]
    if grades:
        total = sum(grades[c.xl_grade_overwrite_cell_id()] for c in rubric.criteria)
        p.text = f"Note : {total}"
        p.style = "Heading 3"

    # Barème et seuils
    set_grade_levels(rubric, hdr_cells)

    # Commentaires
    if grades:
        comment_cell = hdr_cells[-1]
        p = comment_cell.paragraphs[0]
        p.text = "Commentaire"
        p.style = "Heading 3"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_grade_levels(rubric: Rubric, hdr_cells):
    color_schemes = scale_color_schemes(len(rubric.grade_levels))
    for i, label in enumerate(rubric.grade_levels):
        cell = hdr_cells[i + 1]
        if color_schemes:
            set_cell_background(cell, color_schemes[i])
        t_str = rubric.threshold_str(i)
        txt = f"{label}\n({t_str})"
        p = cell.paragraphs[0]
        p.text = txt
        p.style = "Heading 3"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def criterion_has_comments(criterion: Criterion, grades: dict[str, Any] | None) -> bool:
    """
    Vérifie si un critère a des commentaires.
    """
    if not grades:
        return False
    has_comments = False
    id = criterion.xl_comment_cell_id()
    if id in grades and grades[id].strip():
        has_comments = True
    else:
        for indicator in criterion.indicators:
            id = indicator.xl_comment_cell_id()
            if id in grades and grades[id].strip():
                has_comments = True
                break
    return has_comments

def word_to_pdf(docx_path: Path, pdf_path: Path):
    """
    Convertit un document Word (.docx) en PDF sur Windows.
    Utilise l'API COM de Word pour effectuer la conversion.
    """
    if sys.platform != "win32":
        raise NotImplementedError("Word to PDF conversion is only available on Windows.")

    import win32com.client

    # Launch Word
    try:
        # Try to connect to an already running Word instance
        word = win32com.client.GetActiveObject("Word.Application")
        created = False
    except Exception:
        # Start a new one if not found
        word = win32com.client.Dispatch("Word.Application")
        created = True
        word.Visible = False

    try:
        doc = word.Documents.Open(str(docx_path))
        # ExportAsFixedFormat: https://learn.microsoft.com/en-us/office/vba/api/word.document.exportasfixedformat
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,  # wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=0,    # wdExportOptimizeForPrint
            CreateBookmarks=1 # wdExportCreateHeadingBookmarks
        )
        doc.Close(False)
    finally:
        if created:
            # Quitter Word seulement si on l'a créé
            word.Quit()
