import sys
from importlib import resources
from pathlib import Path

import docx
import docx.enum.text
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from c3hm.core.rubric import Rubric


def scale_color_schemes(size: int) -> list[str] | None:
    """
    Retourne une liste de couleurs format Word pour les cellules liées au barème.

    La première couleur est le vert parfait, la dernière est le rouge insuffisant.
    Prend en charge 2 à 5 niveaux de barème. Au-delà, renvoie None.
    """
    perfect_green   = "C8FFC8"  # RGB(200, 255, 200)
    very_good_green = "F0FFB0"  # RGB(240, 255, 176)
    half_way        = "FFF8C2"  # RGB(255, 248, 194)
    minimal_red     = "FFE4C8"  # RGB(255, 228, 200)
    bad_red         = "FFC8C8"  # RGB(255, 200, 200)

    color_schemes = {
        2: [perfect_green, bad_red],
        3: [perfect_green, half_way, bad_red],
        4: [perfect_green, very_good_green, half_way, bad_red],
        5: [perfect_green, very_good_green, half_way, minimal_red, bad_red],
    }
    return color_schemes.get(size)

def repeat_table_header(header):
    """
    Définit la première ligne d'un tableau Word comme en-tête répétée sur chaque page.
    """
    tr = header._tr
    tr_pr = tr.get_or_add_trPr()

    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def set_cell_background(cell, color_hex: str):
    """
    Définit la couleur de fond d'une cellule dans un tableau Word.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')         # shading style
    shd.set(qn('w:color'), 'auto')        # text color (auto = black/white depending on bg)
    shd.set(qn('w:fill'), color_hex)      # background color

    # remove old <w:shd> if any
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    tc_pr.append(shd)

def set_row_borders(row, top=0.5, bottom=0.5):
    """
    Définit les bordures hautes et basses d'une ligne dans un tableau Word.
    """
    def set_border(tc_borders, side, width_pt):
        border = tc_borders.find(qn(f'w:{side}'))
        if border is None:
            border = OxmlElement(f'w:{side}')
            tc_borders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(int(width_pt * 8)))  # Word uses 1/8 pt units
        border.set(qn('w:space'), "0")

    for cell in row.cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        if top:
            set_border(tc_borders, 'top', top)
        if bottom:
            set_border(tc_borders, 'bottom', bottom)

def generate_word_from_rubric(rubric: Rubric, output_path: Path | str) -> None:
    """
    Génère un document Word pour les étudiants à partir d’un grille d’évaluation (Rubric).
    """
    output_path = Path(output_path)

    # forcer l'extension .docx
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    # démarrer Word avec le modèle par défaut
    # Ce dernier contient déjà le formatage et les styles nécessaires
    template = resources.files("c3hm.assets.templates.word").joinpath("rubric-default.docx")
    with resources.as_file(template) as path:
        doc = docx.Document(path)

    # insérer le titre
    heading = doc.add_heading(rubric.title(), level=1)
    heading.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    # insérer la grille
    grid = rubric.grid
    table = doc.add_table(rows=1, cols=grid.nb_columns(), style="Normal Table")

    # Remplir la première ligne avec le barème
    repeat_table_header(table.rows[0])
    set_row_borders(table.rows[0])
    # Première cellule : Total sur X pts
    hdr_cells = table.rows[0].cells
    total_cell = hdr_cells[0]
    p = total_cell.paragraphs[0]
    pts = "pt" if grid.total_score == 1 else "pts"
    p.text = f"Total sur {grid.total_score} {pts}"
    p.style = "Heading 3"
    # Barème et seuils
    color_schemes = scale_color_schemes(len(grid.scale))
    for i, label in enumerate(grid.scale):
        cell = hdr_cells[i + 1]
        if color_schemes:
            set_cell_background(cell, color_schemes[i])
        p = cell.paragraphs[0]
        if i == 0:
            threshold_str = str(grid.thresholds[i])
        else:
            threshold_str = f"{grid.thresholds[i-1] - 1}-{grid.thresholds[i]}"
        p.text = f"{label} [{threshold_str}]"
        p.style = "Heading 3"

    # Remplir le reste avec Critères et indicateurs dans l'ordre
    for criterion in grid.criteria:
        row = table.add_row()
        # Critère
        p = row.cells[0].paragraphs[0]
        pts = "pt" if criterion.weight == 1 else "pts"
        p.text = f"{criterion.name} ({criterion.weight} {pts})"
        p.style = "Heading 3"

        # Indicateurs
        for indicator in criterion.indicators:
            row = table.add_row()
            p = row.cells[0].paragraphs[0]
            run = p.add_run(indicator.name)
            run.style = "Emphasis"

            # Descripteurs alignés avec les niveaux de barème
            for i, descriptor in enumerate(indicator.descriptors):
                row.cells[i + 1].text = descriptor

    # Ajouter une bordure en bas de la dernière ligne
    last_row = table.rows[-1]
    set_row_borders(last_row, top=None, bottom=1.0)

    # enregistrer le fichier
    doc.save(output_path)


def word_to_pdf(docx_path: Path, pdf_path: Path):
    """
    Convertit un document Word (.docx) en PDF sur Windows.
    Utilise l'API COM de Word pour effectuer la conversion.
    """
    if sys.platform != "win32":
        raise NotImplementedError("Word to PDF conversion is only available on Windows.")

    import win32com.client

    # Launch Word
    try:
        # Try to connect to an already running Word instance
        word = win32com.client.GetActiveObject("Word.Application")
        created = False
    except Exception:
        # Start a new one if not found
        word = win32com.client.Dispatch("Word.Application")
        created = True

    word.Visible = False

    try:
        doc = word.Documents.Open(str(docx_path))
        # ExportAsFixedFormat: https://learn.microsoft.com/en-us/office/vba/api/word.document.exportasfixedformat
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,  # wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=0,    # wdExportOptimizeForPrint
            CreateBookmarks=1 # wdExportCreateHeadingBookmarks
        )
        doc.Close(False)
    finally:
        if created:
            # Quitter Word seulement si on l'a créé
            word.Quit()
